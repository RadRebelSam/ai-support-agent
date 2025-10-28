#!/usr/bin/env python3
"""
Simple RAG (Retrieval-Augmented Generation) system without embeddings
Uses direct text matching instead of vector embeddings for document retrieval
"""

# Standard library imports
import os  # File system operations
import re  # Regular expressions for text processing
from typing import List, Dict, Any, Optional  # Type hints for better code documentation
from pathlib import Path  # Cross-platform path handling
import time  # For timing operations and delays

# Third-party imports for web scraping and document processing
import streamlit as st  # Streamlit UI framework
import requests  # HTTP requests for web scraping
from bs4 import BeautifulSoup  # HTML parsing and content extraction
from urllib.parse import urlparse  # URL parsing and validation

# Azure and LangChain imports for AI integration
from config import AzureConfig  # Azure service configuration
from langchain_openai import AzureChatOpenAI  # Azure OpenAI integration
from langchain_core.documents import Document  # Document schema for LangChain

# Document processing libraries (avoiding Windows compatibility issues)
import PyPDF2  # PDF file processing
from docx import Document as DocxDocument  # Word document processing

class SimpleRAGSystem:
    """Simple RAG system using text matching instead of embeddings
    
    This implementation provides a lightweight alternative to vector-based RAG systems
    by using direct text matching and keyword overlap scoring for document retrieval.
    """
    
    def __init__(self):
        # Core data storage
        self.documents = []  # List of processed document chunks
        self.llm = None      # Language model instance (not used in simple mode)
        self.qa_chain = None # Question-answering chain (not used in simple mode)
        
    def load_text_file(self, file_path: str) -> List[Document]:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return [Document(page_content=content, metadata={"source": file_path})]
        except Exception as e:
            st.error(f"Error loading text file {file_path}: {str(e)}")
            return []
    
    def load_pdf_file(self, file_path: str) -> List[Document]:
        """Load PDF file"""
        try:
            documents = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    content = page.extract_text()
                    if content.strip():
                        documents.append(Document(
                            page_content=content,
                            metadata={"source": file_path, "page": page_num}
                        ))
            return documents
        except Exception as e:
            st.error(f"Error loading PDF file {file_path}: {str(e)}")
            return []
    
    def load_docx_file(self, file_path: str) -> List[Document]:
        """Load DOCX file"""
        try:
            doc = DocxDocument(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return [Document(page_content=content, metadata={"source": file_path})]
        except Exception as e:
            st.error(f"Error loading DOCX file {file_path}: {str(e)}")
            return []
    
    def load_url(self, url: str, use_selenium: bool = False) -> List[Document]:
        """Load content from URL with optional JavaScript rendering"""
        try:
            if use_selenium:
                return self._load_url_with_selenium(url)
            else:
                return self._load_url_with_requests(url)
        except Exception as e:
            st.error(f"Error loading URL {url}: {str(e)}")
            return []
    
    def _load_url_with_requests(self, url: str) -> List[Document]:
        """Load content from URL using requests (fast, but no JavaScript)"""
        # Add headers to mimic a real browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text content
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = ' '.join(chunk for chunk in chunks if chunk)
        
        # Get page title
        title = soup.title.string if soup.title else url
        
        # Check if this looks like a JavaScript app
        is_js_app = any([
            "You need to enable JavaScript" in content,
            "enable JavaScript to run this app" in content,
            len(content.strip()) < 100,
            soup.find('div', {'id': 'root'}) is not None,
            soup.find('div', {'id': 'app'}) is not None
        ])
        
        if is_js_app:
            st.warning(f"⚠️ {url} appears to be a JavaScript application. Content may be limited. Try enabling 'Use JavaScript Rendering' for better results.")
        
        return [Document(
            page_content=content,
            metadata={
                "source": url, 
                "type": "url", 
                "title": title,
                "status_code": response.status_code,
                "method": "requests",
                "is_js_app": is_js_app
            }
        )]
    
    def _load_url_with_selenium(self, url: str) -> List[Document]:
        """Load content from URL using Selenium (slower, but handles JavaScript)"""
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            from webdriver_manager.chrome import ChromeDriverManager
            from selenium.webdriver.chrome.service import Service
            import os
            
            # Check if running in cloud environment (Streamlit Cloud)
            is_cloud = os.environ.get('STREAMLIT_SERVER_PORT') is not None
            
            # Setup Chrome options
            chrome_options = Options()
            chrome_options.add_argument("--headless")  # Run in background
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-logging")
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--allow-running-insecure-content")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument("--user-agent=Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
            
            # Additional options for cloud deployment
            if is_cloud:
                chrome_options.add_argument("--remote-debugging-port=9222")
                chrome_options.add_argument("--disable-background-timer-throttling")
                chrome_options.add_argument("--disable-backgrounding-occluded-windows")
                chrome_options.add_argument("--disable-renderer-backgrounding")
                chrome_options.add_argument("--disable-features=TranslateUI")
                chrome_options.add_argument("--disable-ipc-flooding-protection")
            
            # Initialize driver with error handling
            try:
                if is_cloud:
                    # For cloud deployment, try to use system Chrome
                    service = Service()
                else:
                    # For local development, use ChromeDriverManager
                    service = Service(ChromeDriverManager().install())
                
                driver = webdriver.Chrome(service=service, options=chrome_options)
            except Exception as driver_error:
                st.warning(f"ChromeDriver initialization failed: {str(driver_error)}")
                st.info("Falling back to standard HTTP request method...")
                return self._load_url_with_requests(url)
            
            try:
                # Load the page
                driver.get(url)
                
                # Wait for content to load
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                
                # Wait a bit more for JavaScript to render
                time.sleep(3)
                
                # Get page title
                title = driver.title or url
                
                # Get page source after JavaScript execution
                page_source = driver.page_source
                soup = BeautifulSoup(page_source, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style", "nav", "footer", "header"]):
                    script.decompose()
                
                # Extract text content
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                content = ' '.join(chunk for chunk in chunks if chunk)
                
                return [Document(
                    page_content=content,
                    metadata={
                        "source": url, 
                        "type": "url", 
                        "title": title,
                        "status_code": 200,
                        "method": "selenium",
                        "is_js_app": True
                    }
                )]
                
            finally:
                driver.quit()
                
        except ImportError:
            st.error("Selenium not available. Install with: pip install selenium webdriver-manager")
            st.info("Falling back to standard HTTP request method...")
            return self._load_url_with_requests(url)
        except Exception as e:
            st.warning(f"Selenium error loading URL {url}: {str(e)}")
            st.info("Falling back to standard HTTP request method...")
            return self._load_url_with_requests(url)
    
    def is_url(self, path: str) -> bool:
        """Check if input is a URL"""
        try:
            result = urlparse(path)
            return all([result.scheme, result.netloc])
        except:
            return False
    
    def load_documents(self, inputs: List[str], use_js_rendering: bool = False) -> List[Document]:
        """Load documents from files or URLs"""
        documents = []
        
        for input_path in inputs:
            try:
                if self.is_url(input_path):
                    # Handle URL
                    docs = self.load_url(input_path, use_selenium=use_js_rendering)
                    if docs:
                        method = "JavaScript rendering" if use_js_rendering else "standard"
                        st.success(f"🌐 Loaded {len(docs)} documents from URL ({method}): {input_path}")
                else:
                    # Handle file (existing logic)
                    file_extension = Path(input_path).suffix.lower()
                    if file_extension == '.txt':
                        docs = self.load_text_file(input_path)
                    elif file_extension == '.pdf':
                        docs = self.load_pdf_file(input_path)
                    elif file_extension in ['.docx', '.doc']:
                        docs = self.load_docx_file(input_path)
                    else:
                        st.warning(f"Unsupported file type: {file_extension}")
                        continue
                    
                    if docs:
                        st.success(f"📄 Loaded {len(docs)} documents from file: {input_path}")
                
                documents.extend(docs)
                
            except Exception as e:
                st.error(f"Error loading {input_path}: {str(e)}")
                
        return documents
    
    def process_documents(self, documents: List[Document]) -> List[Document]:
        """Process documents by splitting into chunks"""
        if not documents:
            return []
            
        # Simple text splitting
        processed_docs = []
        for doc in documents:
            # Split by paragraphs
            paragraphs = doc.page_content.split('\n\n')
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    processed_docs.append(Document(
                        page_content=paragraph.strip(),
                        metadata={**doc.metadata, "chunk": i}
                    ))
        
        st.info(f"Split documents into {len(processed_docs)} chunks")
        return processed_docs
    
    def create_knowledge_base(self, documents: List[Document]):
        """Create knowledge base from documents"""
        if not documents:
            st.warning("No documents to process")
            return
            
        self.documents = documents
        st.success(f"Created knowledge base with {len(documents)} document chunks")
        
        # Initialize LLM
        try:
            self.llm = AzureChatOpenAI(
                openai_api_key=AzureConfig.OPENAI_API_KEY,
                azure_endpoint=AzureConfig.OPENAI_ENDPOINT,
                openai_api_version=AzureConfig.OPENAI_API_VERSION,
                deployment_name=AzureConfig.OPENAI_DEPLOYMENT,
                temperature=0.7
            )
            st.success("LLM initialized successfully")
        except Exception as e:
            st.error(f"Error initializing LLM: {str(e)}")
            raise
    
    def find_relevant_documents(self, query: str, k: int = 3) -> List[Document]:
        """Find relevant documents using simple text matching"""
        if not self.documents:
            return []
        
        # Simple keyword matching
        query_words = set(re.findall(r'\w+', query.lower()))
        
        scored_docs = []
        for doc in self.documents:
            content_words = set(re.findall(r'\w+', doc.page_content.lower()))
            # Calculate simple overlap score
            overlap = len(query_words.intersection(content_words))
            if overlap > 0:
                scored_docs.append((overlap, doc))
        
        # Sort by score and return top k
        scored_docs.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in scored_docs[:k]]
    
    def query(self, question: str) -> Dict[str, Any]:
        """Query the RAG system"""
        if not self.llm or not self.documents:
            return {
                "answer": "RAG system not initialized. Please upload documents first.",
                "source_documents": []
            }
        
        try:
            # Find relevant documents
            relevant_docs = self.find_relevant_documents(question, k=3)
            
            if not relevant_docs:
                return {
                    "answer": "No relevant information found in the knowledge base.",
                    "source_documents": []
                }
            
            # Create context from relevant documents
            context = "\n\n".join([doc.page_content for doc in relevant_docs])
            
            # Create prompt
            prompt = f"""Based on the following information from the knowledge base, answer the question.

Knowledge Base Information:
{context}

Question: {question}

Please provide a helpful and accurate answer based on the information above. If the information doesn't directly answer the question, say so and provide what information is available."""

            # Get response from LLM
            response = self.llm.invoke(prompt)
            answer = response.content if hasattr(response, 'content') else str(response)
            
            return {
                "answer": answer,
                "source_documents": relevant_docs
            }
            
        except Exception as e:
            return {
                "answer": f"Error querying RAG system: {str(e)}",
                "source_documents": []
            }
    
    def clear_knowledge_base(self):
        """Clear the knowledge base"""
        self.documents = []
        self.llm = None
        st.success("Knowledge base cleared")
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base"""
        if not self.documents:
            return {"total_documents": 0, "status": "Not initialized"}
        
        return {
            "total_documents": len(self.documents),
            "status": "Ready"
        }
